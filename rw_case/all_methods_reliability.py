__author__ = 'zgf'

import datetime

import nibabel as nib
import numpy as np
from docx import Document
from docx.shared import Inches

from configs import *
# from matplot_case.inter_subject_bar import show_barchart

SUBJECT_SESSION_NUM = 7

def dice(volume1, volume2):
    if volume1.shape != volume2.shape:
        raise ValueError("Shape mismatch: volume1 and volume2 must have the same shape.")
    intersection = np.logical_and(volume1, volume2)
    if volume1.sum() + volume2.sum() == 0:
        return 0.0
    else:
        return 2. * intersection.sum() / (volume1.sum() + volume2.sum())

def all_analysis(all_dices):
    means = []
    stds = []
    for i in range(len(all_dices[0])):
        temp_means = []
        temp_stds = []
        for roi_index in range(len(all_dices)):
            temp_means.append(round(all_dices[roi_index][i][all_dices[roi_index][i] >= 0].mean(), 3))
            temp_stds.append(round(all_dices[roi_index][i][all_dices[roi_index][i] >= 0].std(), 3))
        means.append(temp_means)
        stds.append(temp_stds)

    document = Document()
    title = 'Atlas-based Random Walk Methods Analysis '
    document.add_heading('ALL_Subject_Methods Analysis ', 0)

    xlabel = 'ROI'
    ylabel = 'Dice'
    # show_barchart(means, stds, xlabel, ylabel, title, COLOR, METHODS, ROI)

    # document.add_picture(TEMP_IMG_DIR, width=Inches(6.0))
    document.add_paragraph('Mean of Optional Region Size:  ' + str(means), style='ListBullet')
    document.add_paragraph('Std of Optional Region Size:  ' + str(stds), style='ListBullet')
    document.save(ANALYSIS_DIR + 'All_Methods_' + RW_ATLAS_BASED_DOCX_RESULT_FILE)

    print 'all_analysis: ', '  means: ', means
    return means

def inter_subject_analysis(roi_dices, roi_index):
    dices = np.zeros((SUBJECT_SESSION_NUM * (SUBJECT_SESSION_NUM - 1) / 2, roi_dices[0].shape[2], len(roi_dices))).astype(np.float)
    means = []
    stds = []
    for i in range(len(roi_dices)):
        dices[..., i] = roi_dices[i][roi_dices[i] >= 0].reshape((SUBJECT_SESSION_NUM * (SUBJECT_SESSION_NUM - 1) / 2, roi_dices[i].shape[2]))
        means.append(np.average(dices[..., i], axis=0).tolist())
        stds.append(np.std(dices[..., i], axis=0).tolist())

    document = Document()
    title = ROI[roi_index] + '_Methods Analysis '
    document.add_heading(ROI[roi_index] + '_Methods Analysis ', 0)

    xlabel = 'Subject Name'
    ylabel = 'Dice'
    # show_barchart(means, stds, xlabel, ylabel, title, COLOR, METHODS, SUBJECT_NAMES)

    # document.add_picture(TEMP_IMG_DIR, width=Inches(6.0))
    document.add_paragraph('Mean of Optional Region Size:  ' + str(means), style='ListBullet')
    document.add_paragraph('Std of Optional Region Size:  ' + str(stds), style='ListBullet')
    document.save(ANALYSIS_DIR + ROI[roi_index]+ '_Methods_' + RW_ATLAS_BASED_DOCX_RESULT_FILE)
    print 'inter_subject_analysis: ', ' roi_index: ', ROI[roi_index]
    print 'means: ', means

    return means

def manual(roi_index):
    # all_subject_session_path = ANALYSIS_DIR + 'manual/' + ROI[roi_index] +  '_manual.nii.gz'
    # all_subject_session = nib.load(all_subject_session_path).get_data()

    all_subject_session_path = ANALYSIS_DIR + 'manual/' + 'manual.nii.gz'
    all_subject_session = nib.load(all_subject_session_path).get_data()
    all_subject_session[all_subject_session != (roi_index + 1)] = 0
    dices = compute_dice_matrix(all_subject_session)
    return dices

def gss(roi_index):
    ROI = ['r_OFA', 'l_OFA', 'r_pFus', 'l_pFus']
    all_subject_session_path = ANALYSIS_DIR + 'gss/' + ROI[roi_index] + '_gss.nii.gz'
    all_subject_session = nib.load(all_subject_session_path).get_data()
    dices = compute_dice_matrix(all_subject_session)
    return dices

# def ac_srg(roi_index):
#     all_subject_session_path = ANALYSIS_DIR + 'asrg/' + ROI[roi_index] + '_' + AC_OPTIMAL_FILE
#     all_subject_session = nib.load(all_subject_session_path).get_data()
#     dices = compute_dice_matrix(all_subject_session)
#     return dices

def random_walker(roi_index, rw_result_filepath):
    # all_subject_session_path = ANALYSIS_DIR + 'rsrg/' + ROI[roi_index] + '_' + '1000_rsrg.nii.gz'
    # all_subject_session_path = ANALYSIS_DIR + 'rw/' + str(roi_index + 1) + '_rw_result_file.nii.gz'

    # all_subject_session_path = ANALYSIS_DIR + 'rw/' + ROI[roi_index] + '_' + RW_PROB_RESULT_FILE
    # all_subject_session = nib.load(all_subject_session_path).get_data()

    all_subject_session = (nib.load(rw_result_filepath).get_data() == (roi_index + 1)).astype(np.int32)

    dices = compute_dice_matrix(all_subject_session)
    return dices


def compute_dice_matrix(all_subject_session):
    session_length = all_subject_session.shape[3]
    subject_num = session_length / SUBJECT_SESSION_NUM
    # image = all_subject_session.reshape((all_subject_session.shape[0], all_subject_session.shape[1],
    #                                      all_subject_session.shape[2], SUBJECT_SESSION_NUM, subject_num))
    dices = np.zeros((SUBJECT_SESSION_NUM, SUBJECT_SESSION_NUM, subject_num)).astype(np.float) - 1

    for i in range(subject_num):
        temp_image = all_subject_session[..., i * SUBJECT_SESSION_NUM:(i + 1) * SUBJECT_SESSION_NUM]
        for j in range(SUBJECT_SESSION_NUM - 1):
            for k in range(j + 1, SUBJECT_SESSION_NUM):
               dices[j, k, i] = dice(temp_image[..., j] > 0, temp_image[..., k] > 0)

    return dices

if __name__ == "__main__":
    starttime = datetime.datetime.now()

    stats_means = []

    for i in range(20):
        filepath = RW_AGGRAGATOR_RESULT_DATA_DIR + 'staple/rw/top_rank_' + str(i * 10) + '_' + RW_PROB_RESULT_FILE
        # filepath_random = RW_AGGRAGATOR_RESULT_DATA_DIR + 'staple/rw/random_' + str(i * 10) + '_' + RW_PROB_RESULT_FILE
        # filepath = filepath_random
        print 'filepath: ', filepath

        all_dices = []
        for roi_index in range(0, len(ROI)):
        # for roi_index in range(0, 1):
            roi_dices = [manual(roi_index), gss(roi_index), random_walker(roi_index, filepath)]
            inter_subject_analysis(roi_dices, roi_index)
            all_dices.append(roi_dices)
            print 'roi_index:', roi_index
        means = all_analysis(all_dices)
        stats_means.append(means)
        print '--------------------------------------- ', i*10, ' ----------------------------------------'

    print stats_means

    endtime = datetime.datetime.now()
    print 'Time costs: ', (endtime - starttime)
    print "Program end..."






