__author__ = 'zhouguangfu'

import datetime
import multiprocessing

import numpy as np
import nibabel as nib
from docx import Document
from docx.shared import Inches
from scipy.ndimage import morphology
from skimage.segmentation import random_walker

from algorithm.regiongrowing import *
from matplot_case.histogram_features_two import show_date_index_formatter


REGION_SIZE_MAX = 1500

def process_single_subject(subject_image_marker):
    subject_image, markers, subject_index, mask = subject_image_marker
    region_result_RW = np.zeros_like(subject_image)
    region_result_RW[markers[:, 0], markers[:, 1], markers[:, 2]] = 1
    vals = []
    region_size = []
    labels = []

    while region_result_RW.sum() <= REGION_SIZE_MAX:
        #random walker ---------------------------------------------------------------------------------------------
        markers = np.zeros_like(subject_image)
        volume = morphology.binary_dilation(mask).astype(mask.dtype)
        # markers[volume == 0] = -1 #not consider -------------!!!!!!
        markers[volume == 0] = -1
        volume[mask != 0] = 0
        # image[mask == 0, j] = 0
        markers[volume != 0] = 1 #backgroud
        markers[region_result_RW == 1] = 2 #forground
        rw_labels = random_walker(subject_image, markers, beta=10, mode='bf')
        region_result_RW[rw_labels == 2] = 1
        labels.append(copy.copy(region_result_RW))

        # #AC
        # boundary = morphology.binary_dilation(region_result_RW).astype(region_result_RW.dtype)
        # boundary[region_result_RW == 1] = 0
        # boundary_mean = subject_image[boundary != 0].mean()
        # region_mean = subject_image[region_result_RW == 1].mean()
        # val = region_mean - boundary_mean

        #OTSU
        u = subject_image[rw_labels > 0].mean()
        w0 = (rw_labels == 2).sum() * 1. / (rw_labels > 0).sum()
        w1 = 1. - w0
        u0 = subject_image[rw_labels == 2].mean()
        u1 = (u - w0 * u0) / w1
        val = w0 * w1 * (u0 - u1) * (u0 - u1)

        if (len(region_size) > 0 and region_size[-1] == region_result_RW.sum()):
            break
        if region_result_RW.sum() > REGION_SIZE_MAX:
            break
        else:
            vals.append(round(val, 3))
            region_size.append(region_result_RW.sum())

    optional_label = labels[np.array(vals).argmax()]
    print 'subject_index:', subject_index

    return region_size, vals, optional_label

if __name__ == "__main__":
    starttime = datetime.datetime.now()
    # image = nib.load(FOUR_D_DATA_DIR + ALL_SESSION_AVEARGE_FILE)
    image = nib.load(ACTIVATION_DATA_DIR)
    affine = image.get_affine()
    image = image.get_data()

    roi_peak_points = np.load(RESULT_DATA_DIR + RESULT_NPY_FILE)

    lines = []
    for line in open(SUBJECT_ID_DIR):
        if line is not '':
            lines.append(line.rstrip('\r\n'))

    mask = nib.load(ALL_PROB_MASK).get_data()
    document_RW = Document()
    document_RW.add_heading('All prob mask' + ' RW Analysis', 0)

    image_list = []
    for k in range(image.shape[3]):
    # for k in range(4):
        all_prob_markers = np.zeros((27 * len(ROI), 3)).astype(np.int32)
        for i in range(len(ROI)):
            marker = np.array([roi_peak_points[k, i, :]]).astype(np.int32)
            neighbor_element = SpatialNeighbor('connected', mask.shape, 26)
            markers = neighbor_element.compute(marker)
            makers = np.append(markers, marker, axis=0)
            all_prob_markers[27 * i : (i + 1) * 27, :] = makers
        image_list.append((image[..., k], all_prob_markers, k, mask))

    pool = multiprocessing.Pool(processes=multiprocessing.cpu_count() * 2)
    pool_outputs = pool.map(process_single_subject, image_list)
    pool.close()
    pool.join()

    region_result_RW = np.zeros_like(image)
    document_RW = Document()
    document_RW.add_heading('All_prob_mask' + ' RW Analysis', 0)

    for j in range(image.shape[3]):
    # for j in range(4):
        region_size, vals, optional_label = pool_outputs[j]
        region_result_RW[..., j] = optional_label

        title_RW = str(j + 1) + '. ' + lines[j] + ' --  ' + 'All prob mask' +  ' RW Analysis'
        show_date_index_formatter(region_size, vals, 'Region Size', 'RB Value', title_RW, 'g', 'Origin Line', True)
        document_RW.add_heading(str(j + 1) + '. ' + lines[j], 1)
        document_RW.add_paragraph('Region Optimal Size:  ' + str(region_size[np.array(vals).argmax()]), style='ListBullet')
        document_RW.add_paragraph('OTSU Value MAX:  ' +  str(vals[np.array(region_size).argmax()]), style='ListBullet')
        document_RW.add_picture(TEMP_IMG_DIR, width=Inches(4.0))

        print 'j: ', j, '  subject: ', lines[j]
    document_RW.save(RW_RESULT_DATA_DIR + 'all_prob_mask_' + RW_DOCX_RESULT_FILE)
    nib.save(nib.Nifti1Image(region_result_RW, affine), RW_RESULT_DATA_DIR + 'all_prob_mask_' + RW_RESULT_FILE)

    endtime = datetime.datetime.now()
    print 'Time cost: ', (endtime - starttime)
    print "Program end..."
































